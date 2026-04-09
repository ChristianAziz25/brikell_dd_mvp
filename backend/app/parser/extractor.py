"""Document extraction engine — extracts text, structured data, and DD signals.

Supports PDF (PyMuPDF), XLSX/XLS (openpyxl), CSV (pandas), DOCX (python-docx).
"""

import re
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
import pandas as pd
from docx import Document as DocxDocument


MAX_CHARS = 100_000


# ---------------------------------------------------------------------------
# DD signal keyword lists
# ---------------------------------------------------------------------------

_DD_KEYWORDS: dict[str, list[str]] = {
    "contains_rent_roll": [
        "rent roll", "lejefortegnelse", "huslejefortegnelse", "tenant",
        "lejer", "lejemål", "månedlig leje", "monthly rent",
        "lease expiry", "udløb",
    ],
    "contains_floor_areas": [
        "m²", "sqm", "square meter", "areal", "boligareal",
        "erhvervsareal", "kvm", "gla", "nla", "gross leasable",
        "floor area",
    ],
    "contains_energy_data": [
        "energimærke", "energy label", "energiklasse", "kwh",
        "varmeforbrug", "heat consumption", "energy certificate",
        "energiattest",
    ],
    "contains_tenancy_info": [
        "lejekontrakt", "tenancy agreement", "lease agreement",
        "lejebetingelser", "depositum", "deposit", "forudbetalt leje",
        "prepaid rent", "regulering", "indexation",
    ],
    "contains_financial_data": [
        "noi", "driftsresultat", "ebitda", "yield", "afkast",
        "cash flow", "driftsomkostninger", "operating expenses",
        "budget", "regnskab", "annual accounts",
    ],
    "contains_planning_info": [
        "lokalplan", "kommuneplan", "local plan", "plandata",
        "bebyggelsesprocent", "far ", "floor area ratio",
        "anvendelse", "permitted use",
    ],
    "contains_building_specs": [
        "byggeår", "construction year", "opført", "etager", "floors",
        "kælder", "basement", "taglejlighed", "penthouse", "bbr",
        "bygnings", "construction",
    ],
    "contains_ownership_info": [
        "tinglysning", "land registry", "skøde", "title deed",
        "matrikel", "cadastral", "ejerforhold", "ejerlejlighed",
        "owner", "pantebrev", "mortgage", "servitut", "easement",
    ],
    "contains_market_data": [
        "markedsleje", "market rent", "ejendomsværdi", "property value",
        "transaktioner", "transactions", "comparable", "sammenlignelig",
        "prime yield",
    ],
    "contains_opex_data": [
        "grundskyld", "land tax", "ejendomsskat", "forsikring",
        "insurance", "administration", "vicevært", "facility management",
        "vedligeholdelse", "maintenance", "fællesomkostninger",
    ],
}

# Pre-compile a single pattern per signal for speed
_DD_PATTERNS: dict[str, re.Pattern] = {
    key: re.compile("|".join(re.escape(kw) for kw in keywords), re.IGNORECASE)
    for key, keywords in _DD_KEYWORDS.items()
}


def detect_dd_signals(plain_text: str) -> dict[str, bool]:
    """Scan plain_text for DD-relevant keyword groups."""
    return {key: bool(pat.search(plain_text)) for key, pat in _DD_PATTERNS.items()}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_DATE_RE = re.compile(
    r"\b(\d{1,2}[./-]\d{1,2}[./-]\d{2,4}|\d{4}[./-]\d{1,2}[./-]\d{1,2})\b"
)
_NUMBER_RE = re.compile(r"-?\d[\d.,]*\d")
_ADDRESS_RE = re.compile(
    r"\b[A-ZÆØÅa-zæøå][a-zæøå]+(?:vej|gade|allé|vænge|plads|stræde|boulevard|torv)"
    r"\s+\d+[A-Za-z]?(?:\s*,\s*\d{4}\s+[A-ZÆØÅa-zæøå]+)?\b"
)


def _detect_numbers(text: str) -> list[str]:
    return _NUMBER_RE.findall(text)[:200]


def _detect_dates(text: str) -> list[str]:
    return _DATE_RE.findall(text)[:100]


def _detect_addresses(text: str) -> list[str]:
    return _ADDRESS_RE.findall(text)[:50]


def _base_result(filename: str, file_type: str) -> dict:
    return {
        "filename": filename,
        "file_type": file_type,
        "extraction_status": "success",
        "extraction_method": "",
        "page_count": 0,
        "char_count": 0,
        "plain_text": "",
        "structured_data": {
            "tables": [],
            "key_value_pairs": {},
            "numbers_detected": [],
            "dates_detected": [],
            "addresses_detected": [],
        },
        "dd_signals": {k: False for k in _DD_KEYWORDS},
        "error": None,
    }


# ---------------------------------------------------------------------------
# PDF extraction (PyMuPDF / fitz)
# ---------------------------------------------------------------------------

def extract_pdf(file_path: str) -> dict:
    result = _base_result(Path(file_path).name, "pdf")
    result["extraction_method"] = "pymupdf"

    doc = fitz.open(file_path)
    pages_text: list[str] = []
    tables: list[dict] = []
    total_chars = 0

    for i, page in enumerate(doc):
        text = page.get_text("text") or ""
        total_chars += len(text)
        pages_text.append(f"\n\n--- PAGE {i + 1} ---\n\n{text}")

        # Structured blocks
        blocks = page.get_text("blocks") or []
        for block in blocks:
            if block[6] == 0:  # text block
                tables.append({"page": i + 1, "bbox": list(block[:4]), "text": block[4].strip()})

        # Table extraction
        try:
            page_tables = page.find_tables()
            for table in page_tables:
                rows = table.extract()
                if rows:
                    tables.append({"page": i + 1, "rows": rows})
        except Exception:
            pass

    doc.close()

    result["page_count"] = len(pages_text)
    plain = "".join(pages_text)[:MAX_CHARS]
    result["plain_text"] = plain
    result["char_count"] = len(plain)
    result["structured_data"]["tables"] = tables

    if total_chars < 100 and len(pages_text) > 0:
        result["extraction_status"] = "scanned_pdf_low_text"
        result["error"] = (
            "Document appears to be a scanned image. Text extraction limited. "
            "Consider uploading a text-based PDF."
        )

    return result


# ---------------------------------------------------------------------------
# XLSX / XLS extraction (openpyxl)
# ---------------------------------------------------------------------------

def extract_xlsx(file_path: str) -> dict:
    result = _base_result(Path(file_path).name, "xlsx")
    result["extraction_method"] = "openpyxl"

    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_text_parts: list[str] = []
    tables: list[dict] = []
    numbers_detected: list[dict] = []
    dates_detected: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data: list[list[str]] = []
        headers: list[str] = []
        header_found = False

        for row in ws.iter_rows():
            row_vals: list[str] = []
            for cell in row:
                val = cell.value
                if val is None:
                    row_vals.append("")
                else:
                    row_vals.append(str(val))
                    # Track numbers and dates
                    if isinstance(val, (int, float)):
                        numbers_detected.append({
                            "sheet": sheet_name,
                            "cell_ref": cell.coordinate,
                            "value": val,
                        })
                    elif hasattr(val, "strftime"):
                        dates_detected.append({
                            "sheet": sheet_name,
                            "cell_ref": cell.coordinate,
                            "value": str(val),
                        })

            if not header_found and any(v.strip() for v in row_vals):
                headers = row_vals
                header_found = True
            rows_data.append(row_vals)

        sheet_text = f"=== SHEET: {sheet_name} ===\n"
        sheet_text += "\n".join("\t".join(r) for r in rows_data)
        all_text_parts.append(sheet_text)

        tables.append({
            "sheet_name": sheet_name,
            "headers": headers,
            "rows": rows_data,
            "row_count": len(rows_data),
            "col_count": len(headers),
        })

    wb.close()

    plain = "\n\n".join(all_text_parts)[:MAX_CHARS]
    result["plain_text"] = plain
    result["char_count"] = len(plain)
    result["page_count"] = len(wb.sheetnames)
    result["structured_data"]["tables"] = tables
    result["structured_data"]["numbers_detected"] = numbers_detected[:200]
    result["structured_data"]["dates_detected"] = dates_detected[:100]

    return result


# ---------------------------------------------------------------------------
# CSV extraction (pandas)
# ---------------------------------------------------------------------------

def extract_csv(file_path: str) -> dict:
    result = _base_result(Path(file_path).name, "csv")
    result["extraction_method"] = "pandas"

    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="latin-1")

    headers = list(df.columns.astype(str))
    rows = df.astype(str).values.tolist()

    plain = "\t".join(headers) + "\n"
    plain += "\n".join("\t".join(r) for r in rows)
    plain = plain[:MAX_CHARS]

    numbers: list[str] = []
    for col in df.select_dtypes(include="number").columns:
        numbers.extend(str(v) for v in df[col].dropna().tolist())

    result["plain_text"] = plain
    result["char_count"] = len(plain)
    result["page_count"] = 1
    result["structured_data"]["tables"] = [{
        "sheet_name": "csv",
        "headers": headers,
        "rows": rows,
        "row_count": len(rows),
        "col_count": len(headers),
    }]
    result["structured_data"]["numbers_detected"] = numbers[:200]

    return result


# ---------------------------------------------------------------------------
# DOCX extraction (python-docx)
# ---------------------------------------------------------------------------

def extract_docx(file_path: str) -> dict:
    result = _base_result(Path(file_path).name, "docx")
    result["extraction_method"] = "python-docx"

    doc = DocxDocument(file_path)
    paragraphs: list[str] = []
    tables: list[dict] = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            tables.append({
                "headers": rows[0] if rows else [],
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
            })

    plain = "\n".join(paragraphs)
    for t in tables:
        plain += "\n\n[TABLE]\n"
        plain += "\n".join("\t".join(r) for r in t["rows"])
        plain += "\n[/TABLE]"

    plain = plain[:MAX_CHARS]
    result["plain_text"] = plain
    result["char_count"] = len(plain)
    result["page_count"] = len(paragraphs) // 25 + 1
    result["structured_data"]["tables"] = tables

    return result


# ---------------------------------------------------------------------------
# Main router function
# ---------------------------------------------------------------------------

def extract_document(file_path: str, file_type: str, filename: str) -> dict:
    """Extract text and structured data from a document file.

    Returns a dict with plain_text, structured_data, dd_signals, and metadata.
    """
    ext = file_type.lower().lstrip(".")
    try:
        if ext == "pdf":
            result = extract_pdf(file_path)
        elif ext in ("xlsx", "xls"):
            result = extract_xlsx(file_path)
        elif ext == "csv":
            result = extract_csv(file_path)
        elif ext == "docx":
            result = extract_docx(file_path)
        else:
            # Fallback: read as plain text
            text = Path(file_path).read_text(errors="replace")[:MAX_CHARS]
            result = _base_result(filename, ext)
            result["extraction_method"] = "plaintext"
            result["plain_text"] = text
            result["char_count"] = len(text)
            result["page_count"] = 1

        # Override filename with the provided one
        result["filename"] = filename

        # Enrich with detected entities
        plain = result["plain_text"]
        result["structured_data"]["numbers_detected"] = (
            result["structured_data"].get("numbers_detected") or _detect_numbers(plain)
        )
        result["structured_data"]["dates_detected"] = (
            result["structured_data"].get("dates_detected") or _detect_dates(plain)
        )
        result["structured_data"]["addresses_detected"] = _detect_addresses(plain)

        # Run DD signal detection
        result["dd_signals"] = detect_dd_signals(plain)

    except Exception as e:
        result = _base_result(filename, ext)
        result["extraction_status"] = "error"
        result["error"] = str(e)

    return result
