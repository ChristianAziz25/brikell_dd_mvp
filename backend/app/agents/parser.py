"""Document parser — extracts raw text from PDFs, DOCX, CSV, and Excel files.

Also runs the DD extractor and gap analyser after parsing.
No LLM calls. Python-native libraries only.
"""

import json
import sqlite3
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
import pandas as pd

from ..config import settings
from ..parser.extractor import extract_document
from ..parser.dd_gap_analyser import analyse_gaps

MAX_CHARS = 100_000


def _get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(settings.DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    conn.execute("PRAGMA busy_timeout=5000")
    return conn


def _table_to_markdown(table: list[list]) -> str:
    if not table or not table[0]:
        return ""
    headers = table[0]
    rows = table[1:]
    header_line = "| " + " | ".join(str(h) for h in headers) + " |"
    sep_line = "| " + " | ".join("---" for _ in headers) + " |"
    data_lines = ["| " + " | ".join(str(c) for c in row) + " |" for row in rows]
    return "\n".join([header_line, sep_line] + data_lines)


def extract_raw_text(file_path: str, file_type: str) -> str:
    ext = file_type.lower().lstrip(".")
    if ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    elif ext == "csv":
        return _extract_csv(file_path)
    elif ext in ("xlsx", "xls"):
        return _extract_excel(file_path)
    else:
        return Path(file_path).read_text(errors="replace")[:MAX_CHARS]


def _extract_pdf(file_path: str) -> str:
    parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"--- Page {i + 1} ---\n{text}")
            tables = page.extract_tables()
            for table in tables:
                md = _table_to_markdown(table)
                if md:
                    parts.append(f"[TABLE]\n{md}\n[/TABLE]")
    result = "\n\n".join(parts)
    return result[:MAX_CHARS]


def _extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        md = _table_to_markdown(rows)
        if md:
            parts.append(f"[TABLE]\n{md}\n[/TABLE]")
    result = "\n\n".join(parts)
    return result[:MAX_CHARS]


def _extract_csv(file_path: str) -> str:
    df = pd.read_csv(file_path)
    return df.to_markdown(index=False)[:MAX_CHARS]


def _extract_excel(file_path: str) -> str:
    parts: list[str] = []
    xls = pd.ExcelFile(file_path)
    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        parts.append(f"--- Sheet: {sheet} ---\n{df.to_markdown(index=False)}")
    result = "\n\n".join(parts)
    return result[:MAX_CHARS]


def _refresh_gap_report(db: sqlite3.Connection, project_id: str) -> None:
    """Re-run gap analysis for all parsed documents in a project."""
    rows = db.execute(
        "SELECT vault_path, file_type, original_filename FROM documents "
        "WHERE project_id = ? AND parse_status = 'parsed'",
        (project_id,),
    ).fetchall()

    extraction_results = []
    for row in rows:
        result = extract_document(row["vault_path"], row["file_type"], row["original_filename"])
        extraction_results.append(result)

    gap_report = analyse_gaps(project_id, extraction_results)
    db.execute(
        "UPDATE projects SET dd_gap_report = ? WHERE id = ?",
        (json.dumps(gap_report), project_id),
    )
    db.commit()


def parse_document(document_id: str, file_path: str, file_type: str) -> None:
    """Background task: extract raw text, store in DB, then refresh gap analysis."""
    db = _get_db()
    try:
        db.execute(
            "UPDATE documents SET parse_status = 'parsing' WHERE id = ?",
            (document_id,),
        )
        db.commit()
        raw_text = extract_raw_text(file_path, file_type)
        db.execute(
            "UPDATE documents SET raw_text = ?, parse_status = 'parsed' WHERE id = ?",
            (raw_text, document_id),
        )
        db.commit()

        # Get project_id for this document
        row = db.execute(
            "SELECT project_id FROM documents WHERE id = ?", (document_id,)
        ).fetchone()
        if row:
            _refresh_gap_report(db, row["project_id"])

    except Exception as e:
        print(f"Parse error for {document_id}: {e}")
        db.execute(
            "UPDATE documents SET parse_status = 'failed' WHERE id = ?",
            (document_id,),
        )
        db.commit()
    finally:
        db.close()
